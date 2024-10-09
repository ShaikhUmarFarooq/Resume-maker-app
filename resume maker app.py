import tkinter as tk
from tkinter import messagebox
from docx import Document

# Function to generate resume
def generate_resume():
    # Get input values
    name = entry_name.get()
    email = entry_email.get()
    phone = entry_phone.get()
    address = entry_address.get()
    education = text_education.get("1.0", tk.END).strip()
    work_exp = text_work_exp.get("1.0", tk.END).strip()
    skills = text_skills.get("1.0", tk.END).strip()
    
    if not (name and email and phone):
        messagebox.showwarning("Input Error", "Please fill in all required fields.")
        return

    # Create a Word document
    doc = Document()
    
    # Adding header
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    doc.add_paragraph(f"Address: {address}")

    # Education section
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    # Work Experience section
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph(work_exp)

    # Skills section
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)

    # Save the document
    doc_name = f"{name}_Resume.docx"
    doc.save(doc_name)
    messagebox.showinfo("Success", f"Resume saved as {doc_name}")

# GUI
root = tk.Tk()
root.title("Resume Maker")

# Name
tk.Label(root, text="Full Name:").grid(row=0, column=0, padx=10, pady=5)
entry_name = tk.Entry(root, width=40)
entry_name.grid(row=0, column=1)

# Email
tk.Label(root, text="Email:").grid(row=1, column=0, padx=10, pady=5)
entry_email = tk.Entry(root, width=40)
entry_email.grid(row=1, column=1)

# Phone
tk.Label(root, text="Phone:").grid(row=2, column=0, padx=10, pady=5)
entry_phone = tk.Entry(root, width=40)
entry_phone.grid(row=2, column=1)

# Address
tk.Label(root, text="Address:").grid(row=3, column=0, padx=10, pady=5)
entry_address = tk.Entry(root, width=40)
entry_address.grid(row=3, column=1)

# Education
tk.Label(root, text="Education:").grid(row=4, column=0, padx=10, pady=5)
text_education = tk.Text(root, width=30, height=5)
text_education.grid(row=4, column=1)

# Work Experience
tk.Label(root, text="Work Experience:").grid(row=5, column=0, padx=10, pady=5)
text_work_exp = tk.Text(root, width=30, height=5)
text_work_exp.grid(row=5, column=1)

# Skills
tk.Label(root, text="Skills:").grid(row=6, column=0, padx=10, pady=5)
text_skills = tk.Text(root, width=30, height=5)
text_skills.grid(row=6, column=1)

# Generate button
btn_generate = tk.Button(root, text="Generate Resume", command=generate_resume)
btn_generate.grid(row=7, column=1, pady=10)

# Start the GUI
root.mainloop()
